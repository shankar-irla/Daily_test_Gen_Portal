from flask import Flask, render_template, request, send_file
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import ollama
import json
import os

app = Flask(__name__)

FILE_NAME = "daily_coding_memory.json"

generated_content = ""
latest_analysis = ""


def load_memory():
    if os.path.exists(FILE_NAME):
        with open(FILE_NAME, "r") as f:
            return json.load(f)

    return {
        "asked_questions": [],
        "history": []
    }


def save_memory(memory):
    with open(FILE_NAME, "w") as f:
        json.dump(memory, f, indent=4)


def generate_test(memory):

    previous_questions = "\n".join(
        memory["asked_questions"]
    )

    prompt = f"""
    You are a Placement Test Generator.
    
    Previously Asked LeetCode Questions:
    
    {previous_questions}
    
    Generate exactly:
    
    --------------------------------
    
    2 UNIQUE LeetCode Questions
    
    Format:
    Leetcode platform <Question number>
    LC-<number> <problem name> (<difficulty>)
    Problem Description
    Example:
    LC-1 Two Sum (Easy)
    
    --------------------------------
    
    3 UNIQUE Pseudocode MCQs
    
    Format:
    
    Pseudocode Question 1:
    <Question>/<pseudocode>
    
    A) ...
    B) ...
    C) ...
    D) ...
    
    Correct Answer: A
    
    --------------------------------
    
    Rules:
    
    - Never repeat previous LeetCode questions.
    - Coding questions should be Easy-Medium.
    - Pseudocode questions should test logic and tracing.
    - Include correct answers for pseudocode questions.
    - Do not provide explanations.
    - End with: Happy Coding!
    """

    try:

        response = ollama.chat(
            model="llama3.2",
            messages=[
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )

        return response["message"]["content"]

    except Exception as e:

        return f"""
ERROR CONNECTING TO OLLAMA

{str(e)}

Run:

ollama serve
ollama pull llama3.2
"""


def generate_analysis(
    coding_solved,
    pseudo_correct,
    score,
    accuracy
):

    prompt = f"""
Coding Solved: {coding_solved}/2
Pseudocode Correct: {pseudo_correct}/3
Score: {score}/5
Accuracy: {accuracy:.0f}%

Generate:
1. Performance Summary
2. Strengths
3. Weak Areas
4. Motivation
5. Tomorrow Focus
"""

    response = ollama.chat(
        model="llama3.2",
        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ]
    )

    return response["message"]["content"]


@app.route("/")
def home():
    return render_template("index.html")


@app.route("/generate", methods=["POST"])
def generate():

    global generated_content

    memory = load_memory()

    generated_content = generate_test(memory)

    for line in generated_content.split("\n"):

        if line.strip().startswith("LC-"):
            memory["asked_questions"].append(
                line.strip()
            )

    memory["asked_questions"] = list(
        set(memory["asked_questions"])
    )

    save_memory(memory)

    return render_template(
        "paper.html",
        paper=generated_content
    )


@app.route("/submit", methods=["POST"])
def submit():

    global latest_analysis

    coding = int(request.form["coding"])
    pseudo = int(request.form["pseudo"])

    score = coding + pseudo

    accuracy = (score / 5) * 100

    latest_analysis = generate_analysis(
        coding,
        pseudo,
        score,
        accuracy
    )

    return render_template(
        "report.html",
        score=score,
        accuracy=round(accuracy, 2),
        analysis=latest_analysis
    )


@app.route("/download/pdf")
def download_pdf():

    path = "paper.pdf"

    doc = SimpleDocTemplate(path)

    styles = getSampleStyleSheet()

    doc.build([
        Paragraph(
            generated_content.replace(
                "\n",
                "<br/>"
            ),
            styles["BodyText"]
        )
    ])

    return send_file(
        path,
        as_attachment=True
    )


@app.route("/download/docx")
def download_docx():

    path = "paper.docx"

    d = Document()

    d.add_heading(
        "Daily Coding Test",
        level=1
    )

    d.add_paragraph(
        generated_content
    )

    d.save(path)

    return send_file(
        path,
        as_attachment=True
    )


if __name__ == "__main__":

    if not os.path.exists(FILE_NAME):

        with open(FILE_NAME, "w") as f:
            json.dump(
                {
                    "asked_questions": [],
                    "history": []
                },
                f,
                indent=4
            )

    app.run(debug=True)